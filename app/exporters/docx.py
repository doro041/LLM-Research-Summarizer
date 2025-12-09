# app/exporters/docx.py
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.summary import PaperSummary

def to_docx(summary: PaperSummary, output_path: str) -> None:
    """Convert PaperSummary to DOCX format."""
    doc = Document()
    
    # Title
    title = doc.add_heading(summary.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Key Contributions
    doc.add_heading('Key Contributions', level=2)
    for contrib in summary.contributions:
        p = doc.add_paragraph(contrib, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
    
    # Methodology
    doc.add_heading('Methodology', level=2)
    p = doc.add_paragraph(summary.methodology)
    p.paragraph_format.space_after = Pt(12)
    
    # Results
    doc.add_heading('Results', level=2)
    p = doc.add_paragraph(summary.results)
    p.paragraph_format.space_after = Pt(12)
    
    # Limitations
    doc.add_heading('Limitations', level=2)
    for limitation in summary.limitations:
        p = doc.add_paragraph(limitation, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
    
    # Save document
    doc.save(output_path)
